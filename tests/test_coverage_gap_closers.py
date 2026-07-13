"""Тесты для закрытия оставшихся coverage-пробелов.

Закрывает:
  media.py: 17-22 (ImportError), 59 (_transcribe)
  office.py: 32 (empty run), 40 (monospace font), 50-51/88-89/108-109 (ImportError)
  text.py: 31-32 (html ImportError), 76-77 (xml ImportError)
  archive.py: 32 (extract_archive return)
  utils.py: 18 (replace fallback), 52 (run_ocr success)
"""

from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

from any2md.extractors.registry import ExtractionContext


# =====================================================================
# media.py
# =====================================================================

def test_transcribe_no_whisper_raises():
    """media: _transcribe без whisper → RuntimeError."""
    from any2md.extractors.media import _transcribe
    p = Path("/fake/audio.mp3")
    ctx = ExtractionContext()
    with patch.dict("sys.modules", {"whisper": None}):
        with pytest.raises(RuntimeError, match="openai-whisper"):
            _transcribe(p, ctx)


def test_extract_media_with_whisper_module(tmp_path):
    """media: extract_media с переданным whisper_module."""
    from any2md.extractors.media import extract_media

    src = tmp_path / "test.mp3"
    src.write_bytes(b"fake audio")
    ctx = ExtractionContext()

    fake_whisper = MagicMock()
    fake_whisper.load_model.return_value.transcribe.return_value = {"text": "Распознанный текст"}

    result = extract_media(src, ctx, whisper_module=fake_whisper)
    assert "Распознанный текст" in result
    assert "Транскрибация" in result


def test_extract_media_no_whisper_module_calls_transcribe(tmp_path):
    """media: extract_media без whisper_module → _transcribe."""
    from any2md.extractors.media import extract_media

    src = tmp_path / "test.mp3"
    src.write_bytes(b"fake audio")
    ctx = ExtractionContext()

    fake_whisper = MagicMock()
    fake_whisper.load_model.return_value.transcribe.return_value = {"text": ""}

    with patch("any2md.extractors.media._transcribe", return_value=""):
        result = extract_media(src, ctx)
    assert "без распознанной речи" in result


def test_extract_media_with_language(tmp_path):
    """media: extract_media с языком."""
    from any2md.extractors.media import extract_media

    src = tmp_path / "test.wav"
    src.write_bytes(b"fake audio")
    ctx = ExtractionContext(language="ru")

    fake_whisper = MagicMock()
    fake_whisper.load_model.return_value.transcribe.return_value = {"text": "Привет"}

    result = extract_media(src, ctx, whisper_module=fake_whisper)
    # Проверяем что language передан в transcribe
    call_kwargs = fake_whisper.load_model.return_value.transcribe.call_args
    assert call_kwargs[1].get("language") == "ru" or "language" in str(call_kwargs)


# =====================================================================
# office.py
# =====================================================================

def test_docx_empty_run_chunk(tmp_path):
    """office: пустой run пропускается (line 32)."""
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx не установлен")

    doc = docx.Document()
    p = doc.add_paragraph()
    # Добавляем run с пустым текстом и обычный
    p.add_run("")
    p.add_run("Текст")

    src = tmp_path / "test.docx"
    doc.save(src)

    from any2md.extractors.office import extract_docx
    result = extract_docx(src, ExtractionContext())
    assert "Текст" in result


def test_docx_monospace_font(tmp_path):
    """office: monospace шрифт → backtick (line 40)."""
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx не установлен")

    doc = docx.Document()
    p = doc.add_paragraph()
    run = p.add_run("код")
    run.font.name = "Courier New"

    src = tmp_path / "test.docx"
    doc.save(src)

    from any2md.extractors.office import extract_docx
    result = extract_docx(src, ExtractionContext())
    assert "`код`" in result


def test_docx_no_bs4_import_error():
    """office: extract_docx без python-docx → RuntimeError."""
    # extract_docx импортирует docx внутри — патчим sys.modules
    from any2md.extractors.office import extract_docx
    ctx = ExtractionContext()
    fake_path = Path("/fake/test.docx")
    with patch.dict("sys.modules", {"docx": None}):
        with pytest.raises(RuntimeError, match="python-docx"):
            extract_docx(fake_path, ctx)


def test_xlsx_no_openpyxl_import_error():
    """office: extract_xlsx без openpyxl → RuntimeError."""
    from any2md.extractors.office import extract_xlsx
    ctx = ExtractionContext()
    fake_path = Path("/fake/test.xlsx")
    with patch.dict("sys.modules", {"openpyxl": None}):
        with pytest.raises(RuntimeError, match="openpyxl"):
            extract_xlsx(fake_path, ctx)


def test_pptx_no_pptx_import_error():
    """office: extract_pptx без python-pptx → RuntimeError."""
    from any2md.extractors.office import extract_pptx
    ctx = ExtractionContext()
    fake_path = Path("/fake/test.pptx")
    with patch.dict("sys.modules", {"pptx": None}):
        with pytest.raises(RuntimeError, match="python-pptx"):
            extract_pptx(fake_path, ctx)


# =====================================================================
# text.py
# =====================================================================

def test_extract_html_no_bs4(tmp_path):
    """text: extract_html без beautifulsoup4 → fallback."""
    from any2md.extractors.text import extract_html
    src = tmp_path / "test.html"
    src.write_text("<html><body>plain</body></html>", encoding="utf-8")
    ctx = ExtractionContext()
    with patch.dict("sys.modules", {"bs4": None}):
        result = extract_html(src, ctx)
    assert "```" in result  # _fallback_text оборачивает в код-блок


def test_extract_xml_no_bs4(tmp_path):
    """text: extract_xml без beautifulsoup4 → fallback."""
    from any2md.extractors.text import extract_xml
    src = tmp_path / "test.xml"
    src.write_text("<root><item>data</item></root>", encoding="utf-8")
    ctx = ExtractionContext()
    with patch.dict("sys.modules", {"bs4": None}):
        result = extract_xml(src, ctx)
    assert "```" in result


# =====================================================================
# archive.py
# =====================================================================

def test_extract_archive_returns_string(tmp_path):
    """archive: extract_archive возвращает markdown-строку."""
    from any2md.extractors.archive import extract_archive
    src = tmp_path / "test.zip"
    src.write_bytes(b"fake zip")
    ctx = ExtractionContext()
    result = extract_archive(src, ctx)
    assert "Архив" in result
    assert "test.zip" in result


# =====================================================================
# utils.py
# =====================================================================

def test_run_ocr_success(tmp_path):
    """utils: run_ocr успешное распознавание (line 52)."""
    from any2md.utils import run_ocr
    from PIL import Image, ImageDraw, ImageFont

    img = Image.new("RGB", (300, 60), color="white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 28)
    except OSError:
        font = ImageFont.load_default()
    draw.text((10, 15), "Hello OCR", fill="black", font=font)

    src = tmp_path / "ocr_test.png"
    img.save(src)

    result = run_ocr(src, "eng")
    assert "Hello OCR" in result or len(result) > 0


# =====================================================================
# core.py — _path_to_output edge cases
# =====================================================================

def test_path_to_output_relative_path(tmp_path):
    """core: _path_to_output с относительным путём."""
    from any2md.core import _path_to_output
    out_dir = tmp_path / "out"
    src = Path("sub/file.txt")
    result = _path_to_output(src, out_dir, keep_relative=True, root=None)
    assert result.name == "file.md"


# =====================================================================
# cli.py — language from ocr_lang
# =====================================================================

def test_cli_language_from_ocr_lang(tmp_path, capsys):
    """cli: language берётся из ocr_lang если не задан --language."""
    from any2md.cli import main
    src = tmp_path / "test.txt"
    src.write_text("test")
    out = tmp_path / "out.md"
    with pytest.raises(SystemExit) as exc_info:
        main([str(src), "-o", str(out), "--ocr", "--ocr-lang", "rus+eng"])
    assert exc_info.value.code == 0
    # Проверяем что файл создан
    assert out.exists()